"""
cv_processor.py
----------------
CV upload -> parse (PDF/DOCX/TXT) -> refine/optimize (Groq LLM) -> export (.docx)

Flow (yesari kaam garcha):
  1. extract_text_from_cv(file_path)   -> raw text nikalxa CV file bata
  2. optimize_cv_text(raw_text, ...)   -> Groq LLM lai pathayera optimized CV banauxa
  3. save_cv_as_docx(optimized_text)   -> optimized CV lai .docx file ma export garxa
"""

import os
import re
from docx import Document
from docx.shared import Pt, Inches
import pdfplumber
from groq import Groq

GROQ_MODEL = "llama-3.3-70b-versatile"


# ---------------------------------------------------------------------------
# 1. Extraction
# ---------------------------------------------------------------------------
def extract_text_from_cv(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n".join(text_parts)

    elif ext == ".docx":
        doc = Document(file_path)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    else:
        raise ValueError(f"Unsupported CV file format: {ext}. Please upload PDF, DOCX, or TXT.")


# ---------------------------------------------------------------------------
# 2. Optimization via Groq LLM
# ---------------------------------------------------------------------------
def optimize_cv_text(raw_cv_text: str, target_role: str = "", api_key: str = "") -> str:
    """
    Sends the raw CV text to Groq (Llama 3.3 70b) and asks it to return a
    refined, ATS-friendly, well-structured CV in clean Markdown.
    """
    if not api_key:
        api_key = os.environ.get("GROQ_API_KEY", "")
    if not api_key:
        raise RuntimeError("GROQ_API_KEY not set. Please provide it in the sidebar or .env file.")

    client = Groq(api_key=api_key)

    role_line = f"The candidate is targeting roles related to: {target_role}." if target_role else \
                "No specific target role was given — optimize generally for IT/tech roles."

    system_prompt = (
        "You are an expert professional CV/resume writer and ATS (Applicant Tracking System) "
        "optimization specialist. You rewrite CVs to be concise, achievement-oriented, "
        "keyword-optimized for ATS parsing, and well structured. You never invent fake "
        "experience, degrees, or skills that are not implied by the original CV — you only "
        "rephrase, restructure, quantify, and clarify what is already present."
    )

    user_prompt = f"""
Here is a candidate's raw CV text extracted from their uploaded file:

---
{raw_cv_text}
---

{role_line}

Please rewrite this into a polished, professional CV using the following Markdown structure:

# Full Name
Contact line (email | phone | LinkedIn | location) — infer only what's present in the original text.

## Professional Summary
2-3 punchy sentences summarizing strengths and career focus.

## Skills
Grouped bullet list (e.g., Programming Languages, Frameworks/Tools, Soft Skills).

## Experience / Projects
For each entry: Title, Organization/Project name, Dates, then 2-4 bullet points written in strong
action-verb + quantified-impact style (e.g., "Built X that improved Y by Z%").

## Education
Degree, Institution, Dates, relevant details (GPA only if strong).

## Certifications / Achievements (if any were present in the original)

Rules:
- Do NOT fabricate companies, dates, degrees or numbers that aren't implied by the source text.
- Where the original lacks quantification, phrase achievements qualitatively instead of making up fake numbers.
- Keep it to roughly one page worth of content.
- Output ONLY the Markdown CV, no extra commentary before or after.
"""

    response = client.chat.completions.create(
        model=GROQ_MODEL,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.4,
        max_tokens=2000,
    )

    return response.choices[0].message.content.strip()


# ---------------------------------------------------------------------------
# 3. Export optimized CV back to a .docx file
# ---------------------------------------------------------------------------
def save_cv_as_docx(markdown_text: str, output_path: str) -> str:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for section in doc.sections:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)

    lines = markdown_text.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        if stripped.startswith("# "):
            p = doc.add_heading(stripped[2:].strip(), level=0)
        elif stripped.startswith("## "):
            p = doc.add_heading(stripped[3:].strip(), level=1)
        elif stripped.startswith("### "):
            p = doc.add_heading(stripped[4:].strip(), level=2)
        elif stripped.startswith(("- ", "* ")):
            p = doc.add_paragraph(_strip_markdown_bold(stripped[2:].strip()), style="List Bullet")
        else:
            p = doc.add_paragraph(_strip_markdown_bold(stripped))

    doc.save(output_path)
    return output_path


def _strip_markdown_bold(text: str) -> str:
    """Removes simple **bold** markdown markers so plain docx text stays clean."""
    return re.sub(r"\*\*(.*?)\*\*", r"\1", text)